import pytest

from app.ingestion.loader import load_document, load_html, load_text


def test_load_text_basic() -> None:
    content = b"hello world\nfoo bar"
    pages = load_text(content, "test.txt")
    assert len(pages) == 1
    assert pages[0].content == "hello world\nfoo bar"
    assert pages[0].page_number == 1
    assert pages[0].metadata["source"] == "test.txt"


def test_load_markdown() -> None:
    content = b"# Title\n\nSome **bold** text."
    pages = load_document(content, "doc.md", "text/markdown")
    assert len(pages) == 1
    assert "Title" in pages[0].content


def test_load_html_strips_tags() -> None:
    html = b"""
    <html><body>
    <nav>nav links</nav>
    <h1>Main Heading</h1>
    <p>Paragraph content here.</p>
    <script>var x = 1;</script>
    <footer>footer</footer>
    </body></html>
    """
    pages = load_html(html, "page.html")
    assert len(pages) == 1
    text = pages[0].content
    assert "Main Heading" in text
    assert "Paragraph content here." in text
    assert "var x = 1" not in text
    assert "nav links" not in text


def test_load_html_via_dispatch() -> None:
    html = b"<html><body><p>Hello HTML</p></body></html>"
    pages = load_document(html, "test.html", "text/html")
    assert len(pages) == 1
    assert "Hello HTML" in pages[0].content


def test_load_docx_via_dispatch() -> None:
    import io

    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    file_bytes = buf.getvalue()

    pages = load_document(
        file_bytes,
        "test.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert len(pages) == 1
    assert "First paragraph" in pages[0].content
    assert "Second paragraph" in pages[0].content


def test_unsupported_mime_type_raises() -> None:
    with pytest.raises(ValueError, match="unsupported mime type"):
        load_document(b"data", "file.xyz", "application/octet-stream")
